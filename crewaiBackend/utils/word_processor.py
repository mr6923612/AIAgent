# -*- coding: utf-8 -*-
"""
Word文档处理器
用于处理RAG文档目录中的Word文档，提取文字和图片内容
"""

import os
import json
import hashlib
from typing import Dict, List, Tuple, Optional
from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image
import io

class WordDocumentProcessor:
    """Word文档处理器"""
    
    def __init__(self, rag_documents_path: str = "rag_documents"):
        """
        初始化Word文档处理器
        
        Args:
            rag_documents_path: RAG文档目录路径
        """
        self.rag_documents_path = rag_documents_path
        self.processed_docs_file = os.path.join(rag_documents_path, "processed_docs.json")
        self.processed_docs = self._load_processed_docs()
    
    def _load_processed_docs(self) -> Dict:
        """加载已处理的文档记录"""
        try:
            if os.path.exists(self.processed_docs_file):
                with open(self.processed_docs_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
        except Exception as e:
            print(f"加载已处理文档记录失败: {str(e)}")
        return {}
    
    def _save_processed_docs(self):
        """保存已处理的文档记录"""
        try:
            with open(self.processed_docs_file, 'w', encoding='utf-8') as f:
                json.dump(self.processed_docs, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"保存已处理文档记录失败: {str(e)}")
    
    def _calculate_file_hash(self, file_path: str) -> str:
        """计算文件哈希值"""
        try:
            with open(file_path, 'rb') as f:
                file_content = f.read()
                return hashlib.md5(file_content).hexdigest()
        except Exception as e:
            print(f"计算文件哈希失败: {str(e)}")
            return ""
    
    def _is_document_processed(self, file_path: str) -> bool:
        """检查文档是否已经处理过"""
        file_hash = self._calculate_file_hash(file_path)
        return file_hash in self.processed_docs
    
    def _mark_document_processed(self, file_path: str, processing_result: Dict):
        """标记文档为已处理"""
        file_hash = self._calculate_file_hash(file_path)
        self.processed_docs[file_hash] = {
            "file_path": file_path,
            "processed_time": processing_result.get("processed_time", ""),
            "content_count": processing_result.get("content_count", 0),
            "image_count": processing_result.get("image_count", 0)
        }
        self._save_processed_docs()
    
    def process_word_documents(self) -> Dict[str, Dict]:
        """
        处理RAG文档目录中的所有Word文档
        
        Returns:
            处理结果字典
        """
        results = {}
        
        try:
            # 扫描Word文档
            word_files = self._find_word_files()
            
            for word_file in word_files:
                if self._is_document_processed(word_file):
                    print(f"文档 {word_file} 已经处理过，跳过")
                    continue
                
                print(f"正在处理Word文档: {word_file}")
                result = self._process_single_word_document(word_file)
                results[word_file] = result
                
                # 标记为已处理
                self._mark_document_processed(word_file, result)
        
        except Exception as e:
            print(f"处理Word文档失败: {str(e)}")
        
        return results
    
    def _find_word_files(self) -> List[str]:
        """查找Word文档文件"""
        word_files = []
        try:
            for filename in os.listdir(self.rag_documents_path):
                if filename.lower().endswith(('.docx', '.doc')):
                    file_path = os.path.join(self.rag_documents_path, filename)
                    word_files.append(file_path)
        except Exception as e:
            print(f"查找Word文件失败: {str(e)}")
        
        return word_files
    
    def _process_single_word_document(self, file_path: str) -> Dict:
        """处理单个Word文档"""
        try:
            doc = Document(file_path)
            
            # 提取文字内容
            text_content = self._extract_text_content(doc)
            
            # 提取图片
            images = self._extract_images(doc, file_path)
            
            # 关联文字和图片
            content_with_images = self._associate_text_with_images(text_content, images)
            
            # 保存提取的内容
            self._save_extracted_content(file_path, content_with_images)
            
            return {
                "status": "success",
                "processed_time": str(os.path.getmtime(file_path)),
                "content_count": len(content_with_images),
                "image_count": len(images),
                "content": content_with_images
            }
            
        except Exception as e:
            return {
                "status": "error",
                "error": str(e),
                "processed_time": "",
                "content_count": 0,
                "image_count": 0
            }
    
    def _extract_text_content(self, doc: DocumentType) -> List[Dict]:
        """提取文字内容"""
        content = []
        
        # 提取段落
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                content.append({
                    "type": "paragraph",
                    "text": text,
                    "style": paragraph.style.name if paragraph.style else "Normal"
                })
        
        # 提取表格
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            
            if table_data:
                content.append({
                    "type": "table",
                    "data": table_data
                })
        
        return content
    
    def _extract_images(self, doc: DocumentType, file_path: str) -> List[Dict]:
        """提取图片"""
        images = []
        images_dir = os.path.join(self.rag_documents_path, "images")
        
        # 确保图片目录存在
        os.makedirs(images_dir, exist_ok=True)
        
        try:
            # 提取文档中的图片
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_data = rel.target_part.blob
                    
                    # 生成图片文件名
                    image_filename = f"extracted_{len(images) + 1}.png"
                    image_path = os.path.join(images_dir, image_filename)
                    
                    # 保存图片
                    with open(image_path, 'wb') as f:
                        f.write(image_data)
                    
                    # 分析图片
                    try:
                        with Image.open(image_path) as img:
                            images.append({
                                "filename": image_filename,
                                "path": image_path,
                                "width": img.width,
                                "height": img.height,
                                "format": img.format,
                                "size": os.path.getsize(image_path)
                            })
                    except Exception as e:
                        print(f"分析图片失败: {str(e)}")
        
        except Exception as e:
            print(f"提取图片失败: {str(e)}")
        
        return images
    
    def _associate_text_with_images(self, text_content: List[Dict], images: List[Dict]) -> List[Dict]:
        """关联文字和图片内容 - 按产品分组"""
        associated_content = []
        
        # 按空行分割产品
        products = self._split_into_products(text_content)
        
        image_index = 0
        
        for product_id, product_items in enumerate(products, 1):
            # 每个产品作为一个整体
            product_text = []
            product_images = []
            
            for item in product_items:
                if item.get("type") == "paragraph":
                    product_text.append(item.get("text", ""))
                elif item.get("type") == "table":
                    # 将表格转换为文字
                    table_data = item.get("data", [])
                    table_text = " ".join([" ".join(row) for row in table_data])
                    product_text.append(table_text)
            
            # 为每个产品分配图片（如果有的话）
            if image_index < len(images):
                product_images.append(images[image_index])
                image_index += 1
            
            # 创建产品内容项
            if product_text:
                combined_text = "\n".join(product_text)
                associated_content.append({
                    "id": product_id,
                    "text": {
                        "type": "product",
                        "text": combined_text,
                        "style": "Product"
                    },
                    "images": product_images,
                    "product_info": self._extract_product_info(combined_text)
                })
        
        return associated_content
    
    def _split_into_products(self, text_content: List[Dict]) -> List[List[Dict]]:
        """按空行分割产品"""
        products = []
        current_product = []
        
        for item in text_content:
            if item.get("type") == "paragraph":
                text = item.get("text", "").strip()
                
                # 如果遇到空行，开始新产品
                if not text:
                    if current_product:
                        products.append(current_product)
                        current_product = []
                else:
                    current_product.append(item)
            else:
                # 非段落内容（如表格）添加到当前产品
                current_product.append(item)
        
        # 添加最后一个产品
        if current_product:
            products.append(current_product)
        
        return products
    
    def _extract_product_info(self, text: str) -> Dict:
        """提取产品信息"""
        lines = text.split('\n')
        product_info = {
            "name": "",
            "description": "",
            "features": []
        }
        
        if lines:
            # 第一行通常是产品编号+品名
            first_line = lines[0].strip()
            product_info["name"] = first_line
            
            # 其余行是产品介绍
            if len(lines) > 1:
                description_lines = []
                for line in lines[1:]:
                    line = line.strip()
                    if line:
                        description_lines.append(line)
                        # 提取特性（以"-"开头的行）
                        if line.startswith('-') or line.startswith('•'):
                            product_info["features"].append(line[1:].strip())
                
                product_info["description"] = "\n".join(description_lines)
        
        return product_info
    
    def _save_extracted_content(self, file_path: str, content: List[Dict]):
        """保存提取的内容到JSON文件"""
        try:
            base_name = os.path.splitext(os.path.basename(file_path))[0]
            json_filename = f"{base_name}_content.json"
            json_path = os.path.join(self.rag_documents_path, json_filename)
            
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(content, f, ensure_ascii=False, indent=2)
            
            print(f"内容已保存到: {json_path}")
            
        except Exception as e:
            print(f"保存提取内容失败: {str(e)}")
    
    def get_processed_content(self) -> Dict[str, List[Dict]]:
        """获取所有已处理的内容"""
        all_content = {}
        
        try:
            for filename in os.listdir(self.rag_documents_path):
                if filename.endswith('_content.json'):
                    json_path = os.path.join(self.rag_documents_path, filename)
                    with open(json_path, 'r', encoding='utf-8') as f:
                        content = json.load(f)
                        all_content[filename] = content
        except Exception as e:
            print(f"获取已处理内容失败: {str(e)}")
        
        return all_content

# 全局Word文档处理器实例
word_processor = WordDocumentProcessor()
